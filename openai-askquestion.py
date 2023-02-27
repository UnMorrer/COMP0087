import os
import openai
openai.api_key = os.getenv("OPENAI_API_KEY")
import pandas as pd
import numpy as np
import docx
import datetime
from openpyxl import Workbook
def Generate_answers(model,question,number_of_answers):
    #get the distribution of this question
    df=pd.read_excel('C:/Users/tducr/Music/UCL/Term2/COMP0087/project/Data set/1/training_set_rel3.xls')
    data=df.iloc[:,:3]
    data['cnt']=data['essay'].str.split().str.len()
    dist=data[data['essay_set']==question]['cnt']
    length=dist.sample(n=number_of_answers,replace=True).to_numpy()
    temp=np.random.rand(number_of_answers)
    get_test_sample(model=model, length=length, temperature=temp, question=question)

def get_test_sample(model=None,length=[],temperature=[],question=0):
    #test is the input parameters have been properly added
    ls=get_test_prompt()
    if model==None or length==[] or len(length)!=len(temperature) or question not in np.array(ls).T[1] or len(length)!=len(temperature) or max(temperature)>1 or min(temperature)<0:
        return print('Request cannot be performed, wrong argument')
    #generate the answers
    wb=Workbook()
    global ws
    ws=wb.active
    os.chdir(r'C:\Users\tducr\Music\UCL\Term2\COMP0087\project\Data set\AI generate')
    date=datetime.datetime.today().strftime('%Y-%m-%d')
    title=f'AI answers question {question} using the model {model} {date}.xlsx'
    ws.cell(row=1, column=1, value='AI answer')
    ws.cell(row=1, column=2, value='Model')
    ws.cell(row=1, column=3, value='Temperature')
    ws.cell(row=1, column=4, value='Question')
    for idx, l in enumerate(length):
        AI_genarated_answers=generate_AI_genrated_prompt(model,l,temperature[idx],ls[question-1][2])
    #save the answers in am excel file
        #save_AI_answers(AI_genarated_answers[0].choices[0]['text'],question,idx)
        save_AI_answers(AI_genarated_answers,question,idx)
        wb.save(title)

def save_AI_answers(answer,question,i):
    for j in range(3):
        ws.cell(row=2+i,column=j+1).value=answer[j]
    ws.cell(row=2+i,column=4,value=question)

def generate_AI_genrated_prompt(model,length,temperature, question):
    #concatenate the prompt
    quest=concatenate_question(question)
    #call open AI IPA
    #return ['\n\nDear Editor,\n\nAs technology advances, so too does the ubiquity of computers in our lives. While many people believe that this is beneficial for society, I believe that computers have had a negative effect on people.\n\nTo begin with, computers have caused us to become increasingly sedentary. With the ability to work from home and play video games, people are spending more time in front of their screens and less time exercising, which leads to an increase in health problems. Additionally, computers have led to a decrease in face-to-face interaction, which has been linked to higher levels of depression and anxiety.\n\nOn the other hand, computers can provide us with valuable knowledge and skills. They can teach us about faraway places and cultures, as well as give us the ability to communicate with people all over the',model,temperature]
    return [open_ai_API(model,quest,length,temperature).choices[0]['text'],model,temperature]

def open_ai_API(model,prompt,ln,temp):
    return openai.Completion.create(
            model=model,
            prompt=prompt,
            temperature=float(temp),
            max_tokens=int(ln),
            n=1,)
def concatenate_question(ls):
    return '\n'.join(ls)

def get_test_prompt():
    #get the  questions from the .docx files
    path_report=r'C:\Users\tducr\Music\UCL\Term2\COMP0087\project\Data set\1\Essay_Set_Descriptions\Essay_Set_Descriptions'
    list=os.listdir(path_report)#list the docx files
    ls=[]
    #create a list of the docx used for the datasets
    for idx, doc in enumerate(list):
        if doc[:5]=='Essay':
            ls.append([doc,int(doc[doc.find('#')+1])])
    os.chdir(path_report)
    #get the prompt that was used on the student
    for idx, i in enumerate(ls):
        f=open(i[0],'rb')
        doc=docx.Document(f)
        start_source_essay, start_prompt, end_prompt=None, None, None
        for id, j in enumerate(doc.paragraphs):
            try:
                if j.runs[0].text=='Source Essay':
                    start_source_essay=id
                elif j.runs[0].text=='Prompt':
                    start_prompt=id
                elif j.runs[0].text[:17]=='Rubric Guidelines':
                    end_prompt=id
                    break
            except:
                continue
        prompt=[]
        if start_source_essay==None:
            beg=start_prompt
        else:
            beg=start_source_essay
        for j in range(beg+1,end_prompt):
            try:
                prompt.append(doc.paragraphs[j].runs[0].text)
            except:
                continue
        ls[idx].append(prompt)
    return ls

#Generate_answers(model='text-davinci-003',question=2,number_of_answers=618-47)
#Generate_answers(model='text-davinci-003',question=8,number_of_answers=618-254)
#get_test_sample(model='text-ada-001', length=[150,346,567], temperature=[0.4, 0.6, 0.7], question=5)
#pd.read_excel('C:\Users\tducr\Music\UCL\Term2\COMP0087\project\Data set\1\training_set_rel3.xls')